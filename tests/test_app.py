"""Smoke tests for the UIT Portfolio Studio web app.

These run entirely in DEMO mode (no Claude backend, no credentials) and lock in
the behavior rules refined over the project: the SSE turn contract, concept_title
promotion, the docx-skill generation guard, the open-items ([TBD]) surfacing, the
interview-log recovery digest, and the per-project turn lock. Also unit-tests the
Bedrock-aware model resolution, which is otherwise only exercised on the live path.

Run: pip install -r requirements.txt && pytest -q
"""

import json

import docx
import pytest

from conftest import parse_sse


# --------------------------------------------------------------------------
# resolve_model — the Bedrock-aware id resolution (item 4; the live-only path)
# --------------------------------------------------------------------------

def test_resolve_model_explicit_wins(app_mod, monkeypatch):
    monkeypatch.setenv("UIT_CLAUDE_MODEL", "my.custom.profile[1m]")
    monkeypatch.setenv("CLAUDE_CODE_USE_BEDROCK", "1")
    monkeypatch.setenv("ANTHROPIC_MODEL", "global.anthropic.claude-opus-4-8")
    assert app_mod.resolve_model() == "my.custom.profile[1m]"


def test_resolve_model_bedrock_uses_env_profile(app_mod, monkeypatch):
    monkeypatch.delenv("UIT_CLAUDE_MODEL", raising=False)
    monkeypatch.setenv("CLAUDE_CODE_USE_BEDROCK", "1")
    monkeypatch.setenv("ANTHROPIC_MODEL", "global.anthropic.claude-opus-4-8[1m]")
    assert app_mod.resolve_model() == "global.anthropic.claude-opus-4-8[1m]"


def test_resolve_model_bedrock_prefers_fable_env(app_mod, monkeypatch):
    monkeypatch.delenv("UIT_CLAUDE_MODEL", raising=False)
    monkeypatch.delenv("ANTHROPIC_MODEL", raising=False)
    monkeypatch.setenv("CLAUDE_CODE_USE_BEDROCK", "1")
    # Fable is this app's default family: its env profile outranks the Opus one.
    monkeypatch.setenv("ANTHROPIC_DEFAULT_FABLE_MODEL", "global.anthropic.claude-fable-5")
    monkeypatch.setenv("ANTHROPIC_DEFAULT_OPUS_MODEL", "global.anthropic.claude-opus-4-8")
    assert app_mod.resolve_model() == "global.anthropic.claude-fable-5[1m]"


def test_resolve_model_bedrock_appends_1m_to_env_profile(app_mod, monkeypatch):
    monkeypatch.delenv("UIT_CLAUDE_MODEL", raising=False)
    monkeypatch.delenv("ANTHROPIC_MODEL", raising=False)
    monkeypatch.delenv("ANTHROPIC_DEFAULT_FABLE_MODEL", raising=False)
    monkeypatch.setenv("CLAUDE_CODE_USE_BEDROCK", "1")
    # Env profile without the suffix -> [1m] is added (Opus family supports it),
    # and the Opus profile is still honored when no Fable profile is set.
    monkeypatch.setenv("ANTHROPIC_DEFAULT_OPUS_MODEL", "global.anthropic.claude-opus-4-8")
    assert app_mod.resolve_model() == "global.anthropic.claude-opus-4-8[1m]"


def test_resolve_model_bedrock_hardcoded_last_resort(app_mod, monkeypatch):
    for k in ("UIT_CLAUDE_MODEL", "ANTHROPIC_MODEL",
              "ANTHROPIC_DEFAULT_FABLE_MODEL", "ANTHROPIC_DEFAULT_OPUS_MODEL"):
        monkeypatch.delenv(k, raising=False)
    monkeypatch.setenv("CLAUDE_CODE_USE_BEDROCK", "1")
    # Must be a full inference-profile id (the bare alias is rejected by Bedrock).
    assert app_mod.resolve_model() == "global.anthropic.claude-fable-5[1m]"


def test_resolve_model_first_party_short_alias(app_mod, monkeypatch):
    monkeypatch.delenv("UIT_CLAUDE_MODEL", raising=False)
    monkeypatch.delenv("CLAUDE_CODE_USE_BEDROCK", raising=False)
    assert app_mod.resolve_model() == "claude-fable-5[1m]"


def test_resolve_model_opus_family_switch(app_mod, monkeypatch):
    """The per-project downgrade switch resolves through ANTHROPIC_DEFAULT_OPUS_MODEL."""
    monkeypatch.delenv("UIT_CLAUDE_MODEL", raising=False)
    monkeypatch.delenv("ANTHROPIC_MODEL", raising=False)
    monkeypatch.setenv("CLAUDE_CODE_USE_BEDROCK", "1")
    monkeypatch.setenv("ANTHROPIC_DEFAULT_FABLE_MODEL", "global.anthropic.claude-fable-5")
    monkeypatch.setenv("ANTHROPIC_DEFAULT_OPUS_MODEL", "global.anthropic.claude-opus-4-8")
    assert app_mod.resolve_model("opus") == "global.anthropic.claude-opus-4-8[1m]"
    assert app_mod.resolve_model("fable") == "global.anthropic.claude-fable-5[1m]"
    # First-party fallbacks per family
    monkeypatch.delenv("CLAUDE_CODE_USE_BEDROCK", raising=False)
    assert app_mod.resolve_model("opus") == "claude-opus-4-8[1m]"


def test_resolve_model_pins_1m_for_fable_and_opus(app_mod, monkeypatch):
    monkeypatch.delenv("CLAUDE_CODE_USE_BEDROCK", raising=False)
    monkeypatch.setenv("UIT_CLAUDE_MODEL", "claude-fable-5")
    assert app_mod.resolve_model() == "claude-fable-5[1m]"       # Fable gets [1m]
    monkeypatch.setenv("UIT_CLAUDE_MODEL", "claude-opus-4-8[1m]")
    assert app_mod.resolve_model() == "claude-opus-4-8[1m]"      # idempotent


def test_resolve_model_leaves_non_1m_families_alone(app_mod, monkeypatch):
    monkeypatch.delenv("CLAUDE_CODE_USE_BEDROCK", raising=False)
    monkeypatch.setenv("UIT_CLAUDE_MODEL", "claude-sonnet-5")
    assert app_mod.resolve_model() == "claude-sonnet-5"          # Sonnet: no [1m]


# --------------------------------------------------------------------------
# Cost tracking
# --------------------------------------------------------------------------

def test_price_tier_detection(app_mod):
    assert app_mod._price_tier("global.anthropic.claude-opus-4-8[1m]") == "opus"
    assert app_mod._price_tier("claude-sonnet-5") == "sonnet"
    assert app_mod._price_tier("claude-fable-5") == "fable"
    assert app_mod._price_tier("") == "fable"  # default family


def test_cost_tracker_dedupes_by_message_id(app_mod):
    c = app_mod._CostTracker("claude-opus-4-8[1m]")  # $5/$25 per Mtok
    # Same message id updated twice -> latest snapshot wins, not summed.
    c.add_usage("msg_1", {"input_tokens": 1_000_000, "output_tokens": 0})
    assert c.total_usd == pytest.approx(5.0)
    c.add_usage("msg_1", {"input_tokens": 2_000_000, "output_tokens": 0})
    assert c.total_usd == pytest.approx(10.0)  # replaced, not 15
    # A second message adds on.
    c.add_usage("msg_2", {"input_tokens": 0, "output_tokens": 1_000_000})
    assert c.total_usd == pytest.approx(10.0 + 25.0)


def test_cost_tracker_tolerates_null_token_values(app_mod):
    c = app_mod._CostTracker("claude-opus-4-8")
    # A present-but-null field must not crash (get's default only covers absence).
    c.add_usage("msg_1", {"input_tokens": None, "output_tokens": 1_000_000,
                          "cache_read_input_tokens": None})
    assert c.total_usd == pytest.approx(25.0)


def test_cost_tracker_exact_overrides_estimate(app_mod):
    c = app_mod._CostTracker("claude-opus-4-8")
    c.add_usage("msg_1", {"input_tokens": 500_000})
    assert c.total_usd > 0
    c.set_exact(0.4242)
    assert c.total_usd == pytest.approx(0.4242)


# --------------------------------------------------------------------------
# Pure helpers
# --------------------------------------------------------------------------

def test_safe_pid_rejects_traversal(app_mod):
    assert not app_mod._safe_pid("../../etc")
    assert not app_mod._safe_pid("ZZZZ")          # non-hex
    assert not app_mod._safe_pid("abc")           # too short / not a real dir


def test_open_items_lists_every_tbd_with_section(app_mod):
    content = {
        "concept_title": "Enterprise Digital Signage Platform",
        "sponsor": "[TBD — confirm sponsor]",
        "strategic_alignment": "[TBD — name the specific commitment]",
        "define_success": "100% of displays managed",  # no TBD -> excluded
        "risks": "unpatched endpoints",
    }
    items = app_mod._open_items(content)
    sections = {i["section"] for i in items}
    assert "Sponsor" in sections
    assert "Strategic Alignment" in sections
    assert all(i["text"].startswith("[TBD") for i in items)
    assert len(items) == 2  # only the two [TBD] fields


def test_is_concept_brief_detects_template_headings(app_mod, tmp_path):
    d = docx.Document()
    for h in ["CONCEPT TITLE", "CONCEPT DESCRIPTION", "STRATEGIC ALIGNMENT",
              "DEFINE SUCCESS"]:
        d.add_paragraph(h)
    p = tmp_path / "brief.docx"
    d.save(str(p))
    assert app_mod._is_concept_brief(p) is True


def test_is_concept_brief_rejects_plain_doc(app_mod, tmp_path):
    d = docx.Document()
    d.add_paragraph("A strategy memo about GitHub with reviewer comments.")
    p = tmp_path / "memo.docx"
    d.save(str(p))
    assert app_mod._is_concept_brief(p) is False


# --------------------------------------------------------------------------
# API + SSE flow (demo mode)
# --------------------------------------------------------------------------

def test_version_reports_demo_mode(client):
    v = client.get("/api/version").json()
    assert v["demo"] is True                       # CLAUDECODE=1 forces demo
    assert v["docx_skill"] is True                 # installed in this env
    assert v["missing_skills"] == []


def test_create_project_defaults(client):
    p = client.post("/api/projects", json={"mode": "brief"}).json()
    assert p["mode"] == "brief"
    assert p["title"] == "New Concept Brief"
    assert p["version"] == 0
    assert p["has_doc"] is False
    assert p["model_family"] == "fable"  # app default; topbar switch can drop to opus


def test_set_model_family_route(client):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    p = client.post(f"/api/projects/{pid}/model", json={"family": "opus"}).json()
    assert p["model_family"] == "opus"
    # Persisted — a fresh summary reflects it.
    assert client.get(f"/api/projects/{pid}").json()["model_family"] == "opus"
    # Back to the default, and bad input is rejected.
    p = client.post(f"/api/projects/{pid}/model", json={"family": "fable"}).json()
    assert p["model_family"] == "fable"
    r = client.post(f"/api/projects/{pid}/model", json={"family": "haiku"})
    assert r.status_code == 400


def _run_turn(client, pid, text):
    r = client.post(f"/api/projects/{pid}/message", json={"text": text})
    assert r.status_code == 200
    return parse_sse(r.text)


def test_brief_turn_emits_full_contract_and_generates_docx(client, app_mod):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    events = _run_turn(client, pid, "Consolidate campus digital signage onto one platform.")
    kinds = [e for e, _ in events]

    # The SSE turn contract every turn must satisfy.
    assert "project" in kinds
    assert "preview" in kinds
    assert kinds[-1] == "done"
    assert "doc" in kinds  # docx skill present -> a real .docx is produced

    # concept_title promoted to the project title (not the raw first sentence).
    summary = client.get(f"/api/projects/{pid}").json()
    assert summary["title"] == "Enterprise Digital Signage Platform"
    assert summary["version"] >= 1
    assert summary["has_doc"] is True

    # content.json exists on disk and drives the preview.
    content = json.loads((app_mod._pdir(pid) / "out" / "content.json").read_text())
    assert content["concept_title"] == "Enterprise Digital Signage Platform"

    # The generated .docx downloads with the right content type.
    doc_event = next(json.loads(d) for e, d in events if e == "doc")
    dl = client.get(doc_event["url"])
    assert dl.status_code == 200
    assert "wordprocessingml" in dl.headers["content-type"]


def test_preview_surfaces_open_tbd_items(client):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    _run_turn(client, pid, "Digital signage consolidation.")
    pv = client.get(f"/api/projects/{pid}/preview").json()
    # The demo content map leaves [TBD]s in strategic_alignment, sponsor, etc.
    assert pv["open_items"], "expected [TBD] placeholders to surface as open items"
    labels = {i["section"] for i in pv["open_items"]}
    assert "Strategic Alignment" in labels


def test_interview_log_written_for_recovery(client, app_mod):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    _run_turn(client, pid, "My unique answer about patch cadence is 90 days.")
    log = app_mod._pdir(pid) / "out" / "interview-log.md"
    assert log.is_file()
    body = log.read_text()
    assert "My unique answer about patch cadence is 90 days." in body
    assert "## Author" in body and "## Assistant" in body


def test_charter_turn_completes(client):
    pid = client.post("/api/projects", json={"mode": "charter"}).json()["id"]
    events = _run_turn(client, pid, "Turn my reviewed brief into a charter.")
    assert [e for e, _ in events][-1] == "done"


def test_turn_lock_returns_409_when_busy(client, app_mod):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    # Simulate an in-flight turn by claiming the guard directly.
    app_mod._active_turns.add(pid)
    try:
        r = client.post(f"/api/projects/{pid}/message", json={"text": "hi"})
        assert r.status_code == 409
        assert "already being generated" in r.json()["error"]
    finally:
        app_mod._active_turns.discard(pid)
    # Guard released -> the next turn works.
    assert _run_turn(client, pid, "ok now")[-1][0] == "done"


def test_turn_lock_released_after_normal_turn(client, app_mod):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    _run_turn(client, pid, "hello")
    assert pid not in app_mod._active_turns  # released in the generator's finally


def test_message_unknown_project_404(client):
    r = client.post("/api/projects/deadbeef/message", json={"text": "hi"})
    assert r.status_code == 404


# --------------------------------------------------------------------------
# Per-browser session isolation
# --------------------------------------------------------------------------

def test_projects_isolated_between_sessions(client, client2, app_mod):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]

    # client2 is a different browser (different session cookie) — it must not
    # see the project in its list, nor be able to fetch/act on it directly.
    assert pid not in [p["id"] for p in client2.get("/api/projects").json()]
    assert client2.get(f"/api/projects/{pid}").status_code == 404
    assert client2.get(f"/api/projects/{pid}/preview").status_code == 404
    assert client2.post(f"/api/projects/{pid}/message", json={"text": "hi"}).status_code == 404
    assert client2.post(f"/api/projects/{pid}/mode",
                        json={"mode": "charter"}).status_code == 404

    # A delete from the wrong session is a silent no-op — the project survives
    # and its owner can still reach it.
    client2.delete(f"/api/projects/{pid}")
    assert client.get(f"/api/projects/{pid}").status_code == 200

    # The owning session sees and can use it normally.
    assert pid in [p["id"] for p in client.get("/api/projects").json()]
    assert client.get(f"/api/projects/{pid}").status_code == 200


def test_download_respects_ownership(client, client2, app_mod):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    events = _run_turn(client, pid, "Digital signage consolidation.")
    doc_event = next(json.loads(d) for e, d in events if e == "doc")
    # The owner can download; a different session hitting the same URL cannot,
    # even though it knows the exact project id and filename.
    assert client.get(doc_event["url"]).status_code == 200
    assert client2.get(doc_event["url"]).status_code == 404


def test_legacy_project_without_owner_visible_to_every_session(client, client2, app_mod):
    # Simulate a project created before per-session ownership existed (no
    # "owner" key in meta.json) — it must stay visible to every session, same
    # as the app's behavior before this feature, so nothing already on disk
    # becomes inaccessible.
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    meta = app_mod._load_meta(pid)
    del meta["owner"]
    app_mod._save_meta(pid, meta)

    assert pid in [p["id"] for p in client.get("/api/projects").json()]
    assert pid in [p["id"] for p in client2.get("/api/projects").json()]
    assert client2.get(f"/api/projects/{pid}").status_code == 200


def test_upload_detects_brief_without_switching_mode(client, app_mod, tmp_path):
    pid = client.post("/api/projects", json={"mode": "brief"}).json()["id"]
    d = docx.Document()
    for h in ["CONCEPT TITLE", "CONCEPT DESCRIPTION", "STRATEGIC ALIGNMENT",
              "PROPOSED SOLUTION"]:
        d.add_paragraph(h)
    f = tmp_path / "prior_brief.docx"
    d.save(str(f))
    with open(f, "rb") as fh:
        r = client.post(f"/api/projects/{pid}/upload",
                        files={"file": ("prior_brief.docx", fh,
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    body = r.json()
    assert body["is_brief"] is True
    # Reporting only — the mode must NOT have auto-switched to charter.
    assert body["mode"] == "brief"
