#!/usr/bin/env python3
"""Fill the UIT Concept/Idea Brief template from a JSON content map.

The brief template (assets/ConceptBrief_Template.docx) is a single Word table
laid out as a 6-column grid. Each section is a **gray label cell** on the left
(`fill=E7E6E6`, holding the heading and its italic prompt) paired with a **white
answer cell** to its right. The author's content belongs in the WHITE answer
cell — NEVER appended into the gray label cell.

Two placements:
  "right"  -> write into the white answer cell immediately to the right of the
              gray label cell, in the same row. Used for the key/value rows
              (Concept Title, Requestor/Sponsor, dates) and the main narrative
              sections (Concept Description, Strategic Alignment, Outcomes, …).
  "below"  -> a few sections (Key Requirements, Teams & Skills, Additional Notes)
              are a full-width gray heading row followed by a full-width white
              row; their content goes in the white row beneath.

This reproduces a hand-filled template, preserving the original styling. It only
writes the keys you provide — anything omitted is left blank for the human
(we never invent facts).

Usage:
    python fill_brief.py --template <template.docx> --out <out.docx> --content content.json

content.json keys (all optional):
    concept_title, requestor, requestor_title_unit, sponsor, sponsor_title_unit,
    request_date, concept_description, strategic_alignment, outcomes,
    benefits, risks, define_success, proposed_solution, desired_completion_date,
    prioritization_drivers, key_requirements, teams_and_skills, additional_notes
"""
import argparse
import json
import sys

try:
    import docx
except ImportError:
    sys.exit("python-docx is required: pip install python-docx")

# (json_key, anchor_text_the_label_cell_starts_with, placement, occurrence_index)
#   "right" -> white answer cell to the RIGHT of the matched gray label cell
#   "below" -> white answer row BENEATH a full-width gray heading row
SECTIONS = [
    ("concept_title",           "CONCEPT TITLE",            "right", 0),
    ("requestor",               "REQUESTOR",                "right", 0),
    ("sponsor",                 "SPONSOR",                  "right", 0),
    ("requestor_title_unit",    "Title/Unit",               "right", 0),
    ("sponsor_title_unit",      "Title/Unit",               "right", 1),
    ("request_date",            "Request Date",             "right", 0),
    ("concept_description",     "CONCEPT DESCRIPTION",      "right", 0),
    ("strategic_alignment",     "STRATEGIC ALIGNMENT",      "right", 0),
    ("outcomes",                "OUTCOMES",                 "right", 0),
    ("benefits",                "BENEFITS",                 "right", 0),
    ("risks",                   "RISKS",                    "right", 0),
    ("define_success",          "DEFINE SUCCESS",           "right", 0),
    ("proposed_solution",       "PROPOSED SOLUTION",        "right", 0),
    ("desired_completion_date", "DESIRED COMPLETION DATE",  "right", 0),
    ("prioritization_drivers",  "PRIORITIZATION DRIVERS",   "right", 0),
    ("key_requirements",        "KEY REQUIREMENTS",         "below", 0),
    ("teams_and_skills",        "TEAMS and SKILLS",         "below", 0),
    ("additional_notes",        "ADDITIONAL NOTES",         "below", 0),
]


def distinct_cells(row):
    """Cells of a row in grid order, skipping merge repeats."""
    out, seen = [], set()
    for cell in row.cells:
        if cell._tc in seen:
            continue
        seen.add(cell._tc)
        out.append(cell)
    return out


def find_anchor(table, anchor, occurrence):
    """Return (row_index, distinct_index, distinct_cells) for the Nth distinct
    cell whose text starts with anchor, or (None, None, None)."""
    hits = 0
    for ri, row in enumerate(table.rows):
        dc = distinct_cells(row)
        for di, cell in enumerate(dc):
            if cell.text.strip().startswith(anchor):
                if hits == occurrence:
                    return ri, di, dc
                hits += 1
    return None, None, None


def write_cell(cell, text):
    """Write text into a cell, reusing its first empty paragraph then adding more.
    Does not disturb the cell's existing styling."""
    first = True
    for line in str(text).split("\n"):
        if first and cell.paragraphs and not cell.paragraphs[0].text.strip():
            cell.paragraphs[0].add_run(line)
            first = False
        else:
            cell.add_paragraph(line)
            first = False
    return True


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--content", required=True, help="path to content JSON")
    args = ap.parse_args()

    with open(args.content) as f:
        content = json.load(f)

    doc = docx.Document(args.template)
    table = doc.tables[0]

    filled, missing = [], []
    for key, anchor, placement, occ in SECTIONS:
        val = content.get(key)
        if val is None or str(val).strip() == "":
            missing.append(key)
            continue
        ri, di, dc = find_anchor(table, anchor, occ)
        if ri is None:
            print(f"WARNING: anchor not found for '{key}' ({anchor!r})", file=sys.stderr)
            continue
        if placement == "right":
            # The white answer cell is the next distinct cell in the same row.
            if di + 1 < len(dc):
                write_cell(dc[di + 1], val)
                filled.append(key)
            else:
                print(f"WARNING: no answer cell right of '{key}' ({anchor!r})", file=sys.stderr)
        else:  # "below" — white row beneath the full-width heading
            if ri + 1 < len(table.rows):
                write_cell(distinct_cells(table.rows[ri + 1])[0], val)
                filled.append(key)
            else:
                print(f"WARNING: no row below '{key}' ({anchor!r})", file=sys.stderr)

    doc.save(args.out)
    print(f"Wrote {args.out}")
    print(f"Filled ({len(filled)}): {', '.join(filled)}")
    if missing:
        print(f"Left blank ({len(missing)}): {', '.join(missing)}")


if __name__ == "__main__":
    main()
