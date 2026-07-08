#!/usr/bin/env python3
"""UIT Portfolio Studio — a chat + wizard web app for writing OSU UIT
Concept Briefs and Project Charters, powered by Claude Code + the
``uit-portfolio`` skill.

Architecture (modeled on dirkpetersen/codecheck):
- FastAPI single process serves the SPA (static/index.html) and JSON/SSE APIs.
- Work is organized into persistent **projects** (the left-hand list). Each
  project is a directory holding uploads, generated .docx versions, the chat
  transcript, and a live content map that drives the document preview.
- A chat message runs Claude Code in the project directory with
  ``--output-format stream-json``; the skill interviews the author and writes a
  .docx into ``out/`` plus an ``out/content.json`` describing the sections.
- Three Claude tiers, auto-detected (same idea as codecheck):
    1. Claude Code CLI (preferred). Point it at Bedrock with
       CLAUDE_CODE_USE_BEDROCK=1 in the environment for an OSU-approved boundary.
    2. (CLI + Bedrock is the production path — Bedrock is configured via the
       CLI's own env, so there is no separate SDK path here.)
    3. DEMO mode — used when no CLI is available (or inside a Claude Code
       session, where nested invocation is blocked). It streams a scripted
       interview and generates a real .docx via the skill's fill_brief.py, so
       the whole app is runnable and testable without any backend credentials.
"""

import asyncio
import json
import os
import re
import shutil
import subprocess
import time
import uuid
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

from fastapi import FastAPI, Request, UploadFile, File
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles

REPO_ROOT = Path(__file__).parent
SKILL_DIR = REPO_ROOT / ".claude" / "skills" / "uit-portfolio"
FILL_BRIEF = SKILL_DIR / "scripts" / "fill_brief.py"
BRIEF_TEMPLATE = SKILL_DIR / "assets" / "ConceptBrief_Template.docx"
CHARTER_TEMPLATE = SKILL_DIR / "assets" / "UIT_Charter_Scope_Statement_Template_v2.docx"

# Anthropic helper skills the app relies on — ALL are HARD REQUIREMENTS, but in
# two flavors:
#  • FILE_SKILLS are filesystem skills installed under ~/.claude/skills/. A Word
#    template must never be filled without `docx`; uploads must be read with their
#    matching skill (`pdf`/`pptx`/`xlsx`) rather than ad-hoc text extraction. These
#    are existence-checked and symlinked into every project dir.
#  • `claude-api` is a BUILT-IN Claude Code skill (no file on disk), so it can't be
#    existence-checked or symlinked — it's enforced via the PREAMBLE, which tells
#    the model to consult it for any Claude/Anthropic API guidance.
GLOBAL_SKILLS_DIR = Path.home() / ".claude" / "skills"
FILE_SKILLS = ["docx", "pdf", "pptx", "xlsx"]
HELPER_SKILLS = FILE_SKILLS  # symlinked into each project's .claude/skills/


def _skill_installed(name: str) -> bool:
    return (GLOBAL_SKILLS_DIR / name / "SKILL.md").is_file()


def missing_skills() -> list[str]:
    """File-backed skills that are not installed — the app is degraded if non-empty.
    (`claude-api` is built into the CLI, so it is enforced via the PREAMBLE, not here.)"""
    return [s for s in FILE_SKILLS if not _skill_installed(s)]


def docx_skill_available() -> bool:
    """True only if the Anthropic docx skill is installed and loadable."""
    return _skill_installed("docx")

DATA_DIR = Path(os.environ.get("UIT_DATA_DIR", str(Path.home() / ".uit-portfolio")))
PROJECTS_DIR = DATA_DIR / "projects"
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)

app = FastAPI(title="UIT Portfolio Studio")
app.mount("/static", StaticFiles(directory=str(REPO_ROOT / "static")), name="static")

try:
    _GIT_COMMIT = subprocess.check_output(
        ["git", "rev-parse", "--short", "HEAD"], cwd=REPO_ROOT, text=True
    ).strip()
except Exception:
    _GIT_COMMIT = ""

# Canonical section order for the document preview (content.json key -> label).
SECTION_LABELS = [
    ("concept_description", "Concept Description"),
    ("strategic_alignment", "Strategic Alignment"),
    ("outcomes", "Outcomes / Key Results"),
    ("benefits", "Benefits (cost of inaction)"),
    ("risks", "Risks"),
    ("define_success", "Define Success"),
    ("proposed_solution", "Proposed Solution"),
    ("desired_completion_date", "Desired Completion Date"),
    ("prioritization_drivers", "Prioritization Drivers"),
    ("key_requirements", "Key Requirements"),
    ("teams_and_skills", "Teams & Skills Needed"),
    ("additional_notes", "Additional Notes"),
]
FIELD_LABELS = [
    ("requestor", "Requestor"),
    ("requestor_title_unit", "Title / Unit"),
    ("sponsor", "Sponsor"),
    ("sponsor_title_unit", "Sponsor Title / Unit"),
    ("request_date", "Request date"),
]

PREAMBLE = """\
You are the engine behind the UIT Portfolio Studio web app.

BE BRIEF. The author reads your text in a small chat bubble, not a report.
- Summarize what you did or found in 1–3 short sentences, then stop.
- Do NOT narrate your steps ("Let me read X", "Now I'll check the template") —
  just do the work silently and report the result.
- Do NOT restate the author's document back to them, and do NOT write
  multi-paragraph analysis. Save all that depth for the .docx you generate.
- No filler openers like "Certainly". Get to the point.

Working directory rules (IMPORTANT, override conflicting guidance):
- The uit-portfolio skill lives at ./.claude/skills/uit-portfolio/ relative to
  your working directory. Its templates, references, and scripts are in that
  folder — NOT in the working directory itself, so do not go looking for them
  there. The exact paths are:
    • Brief template:   ./.claude/skills/uit-portfolio/assets/ConceptBrief_Template.docx
    • Charter template: ./.claude/skills/uit-portfolio/assets/UIT_Charter_Scope_Statement_Template_v2.docx
    • Fill script:      ./.claude/skills/uit-portfolio/scripts/fill_brief.py
    • References:       ./.claude/skills/uit-portfolio/references/*.md
- Any files the author uploaded are in ./uploads/ — read them first.
- When you generate the document, write the .docx into ./out/. Name a brief
  "<short-name>_Concept_Brief.docx" and a charter "<short-name>_Charter.docx".
- ALSO write the section content map you used to ./out/content.json. Valid keys:
  concept_title, requestor, requestor_title_unit, sponsor, sponsor_title_unit,
  request_date, concept_description, strategic_alignment, outcomes, benefits,
  risks, define_success, proposed_solution, desired_completion_date,
  prioritization_drivers, key_requirements, teams_and_skills, additional_notes.
  This drives the live preview pane. Write it every time you change the document.
- Write `concept_title` EARLY. As soon as you understand what the project is —
  on your FIRST pass, before asking questions or drafting the full document —
  write at least {"concept_title": "<short recognisable name>"} to
  ./out/content.json. The app uses it to title the project, so the sidebar stops
  saying "New Concept Brief" once you have enough to name it. Refine it later if
  the project's focus shifts.
- For a CHARTER, fill the charter template (path above) with the python-docx
  library (no fill script exists for charters) — that is the tool the `docx`
  skill itself prescribes, so follow that skill's guidance and never hand-edit
  raw OOXML. Use the fill_below() helper pattern in
  ./.claude/skills/uit-portfolio/references/charter-fields.md.
  Name the output file "<short-name>_Charter.docx" and write it to ./out/.
- The Anthropic skills `docx`, `pdf`, `pptx`, `xlsx`, and `claude-api` are loaded.
  ALWAYS read an uploaded file with its matching skill — never ad-hoc text
  extraction:
    • .docx → `docx` skill        • .pdf  → `pdf` skill
    • .pptx → `pptx` skill        • .xlsx → `xlsx` skill
  Follow the `docx` skill's guidance for writing/filling Word documents too
  (fill_brief.py and the python-docx charter path both conform to it). Do NOT
  hand-edit raw OOXML. Only fall back to plain extraction if the matching skill
  is genuinely unavailable. If the `docx` skill is not available, do NOT fill or create the
  Word template — tell the author and stop. If anything you say touches the Claude
  or Anthropic API (model ids, pricing, params), consult the `claude-api` skill
  rather than answering from memory.
- Give the document a short, descriptive `concept_title` — a few words a person
  would recognise (e.g. "Enterprise Digital Signage Platform"), NOT a sentence
  copied from the author's request.

Skill rules — these come from the uit-portfolio skill and override defaults:

1. NEVER invent facts. Funding numbers, dates, names, vendor names, protocols,
   integration methods, architectures — these come from the author only.
   You may offer framing and structure ("here's how to phrase this outcome")
   but never suggest specific technical or factual content the author hasn't
   given you.

   PREFER ASKING OVER PLACEHOLDERING. A "[TBD]" is a last resort, not a default.
   - If it's a simple thing the AUTHOR can answer — desired completion date,
     rough timeline, target scale, which units are in scope — just ASK it in
     your questions block. Do NOT drop a "[TBD]" for something they could tell
     you in one line. (A blank "desired completion date" the author was never
     asked is exactly the kind of gap that should have been a question.)
   - Use "[TBD — confirm with <who>]" only when the answer genuinely needs
     someone else (a sponsor's committed funding figure, a HECVAT result) or is
     truly unknowable right now — and even then, ask the author first if they
     might know.

2. NEVER decide. No "approved," no pass/fail, no priority ranking. That includes
   the Pipeline review outcome: reviewer comments in a brief only prove it was
   *reviewed* — never write "approved" (or any outcome) into a charter or summary
   unless the author confirmed it; say "reviewed with comments" or mark it
   "[TBD — confirm Pipeline outcome with the author]".

3. TIER DISCIPLINE — which questions you ask depends on the mode:
   - BRIEF MODE: ask only Tier 1 questions (lightweight, directional).
     Do NOT ask about SLAs, HECVAT, itemized funding, ZTA, data classification,
     backup/DR, per-team ROM, or performance/QoS — that detail belongs in the
     charter, and asking for it now turns a quick intake into a bureaucratic wall.
   - CHARTER MODE: Tier 1 is already covered by the brief. Focus on Tier 2:
     scope in/out, data classification (name the regime: FERPA/HIPAA/GLBA/
     PCI/CUI), stakeholder ROM, funding one-time vs. recurring, post-go-live
     ownership, HECVAT yes/no, ZTA fit, SOC/SIEM telemetry, WCAG, vendor data
     destruction, OCM/comms plan, build-vs-buy, deprecation dependencies.

4. QUESTIONS — ask about 5 PER BATCH (per turn), in their own block, each short.
   ≈5 is the batch size, NOT a total cap. Over the conversation you may ask many
   batches — dozens of questions in total is fine if each genuinely strengthens
   the proposal. Don't pad a batch to 5 if you have fewer real questions, and
   don't stop early just to be brief: keep going as long as it adds value, while
   letting the author end it whenever they want (see rule 5).
   On your FIRST batch, tell the author how to respond, in one short line:
   they can answer by number (e.g. "1. next quarter  2. ~5 colleges  3. not
   sure"), answer only the ones they want, and download/edit the draft anytime.
   - ORDER BY RELEVANCE — ask the MOST relevant questions first. Lead with the
     gaps that matter most for THIS specific brief: the weaknesses a reviewer
     would flag hardest and the answers that most strengthen the proposal. Tackle
     the biggest gaps in the early batches; save minor or tangential details for
     later. Never open with random or low-value questions.
   - Put commentary in prose; put the questions in a fenced block labelled
     `questions`. Never bury a question inside a paragraph of commentary.
   - ONE atomic question per line. NEVER put two questions on one line. If a line
     joins two asks with "and", or has two question marks, SPLIT it into two
     lines. Offering options *for a single question* is fine; cramming separate
     questions is not.
       Good:  Is the funding committed, requested, or still on the radar?
       Bad:   Which funding model, and is it committed, and what user scale?
   - Keep each line to one short sentence. Up to ~5 lines. Example:
     ```questions
     What's driving the timing — a compliance date, a dependency, or leadership?
     Is this adding something new, or replacing an existing system?
     Roughly what user scale are we sizing for?
     Is the funding committed, requested, or still on the radar?
     Who is the single accountable decision-maker?
     ```
   - Do NOT number the lines yourself — the app numbers them so the author can
     answer by number. Just one plain question per line.
   - If you have nothing to ask this turn, omit the block entirely.

5. KEEP ASKING REAL QUESTIONS — never ask permission to continue. The author
   decides when to stop; your job is just to supply substance.
   - After EACH batch, fold the answers in — update ./out/content.json and
     regenerate the .docx — so a current, improved draft is always downloadable.
   - Then simply ask the NEXT batch of substantive questions. Do NOT end a turn
     with a meta-question about whether to continue or download. For example,
     NEVER write things like "Want another batch of questions to tighten anything,
     or are you good to download as-is?" The draft is always available to download
     (the UI makes that clear), so you never prompt about downloading or about
     whether to keep going — you just ask the next useful questions.
   - Stop only when you genuinely run out of useful questions, or when the author
     says "I don't know" / "that's enough" / stops answering. Then leave the
     remaining gaps as `[TBD]` placeholders and hand them the latest draft.

6. BRIEF TIER 1 QUESTIONS TO PRIORITISE (pick the ones that apply — about 5/turn):
   - Problem clarity: what's the actual problem and who feels it?
   - Desired outcomes: what measurably changes if this works? (result, not tasks)
   - Strategic alignment: which specific Prosperity Widely Shared goal? OFFER A
     PICK-LIST — read ./.claude/skills/uit-portfolio/references/strategic-plan.md
     and present its THREE GOALS as a pick-list (they fit one questions block),
     rather than asking the author to recall one. Take their pick verbatim, then
     sharpen it with the matching action or 2030 target from that file and ask HOW
     the work advances it; never invent the goal or the connection.
   - Directional scope: adding new, or replacing/retiring existing? Which units?
   - Who's impacted / who to loop in early (flag OIS/security if relevant)?
   - Major dependencies: does this depend on or enable other work?
   - Rough resource sense (ROM): order-of-magnitude people/tech/budget; any
     funding identified yet, or still "on the radar"?
   - Prioritization driver: what's driving the timing — compliance date,
     dependency, leadership ask? "ASAP" without a reason is weak.
   - Desired completion / timeline: when does the author want this done — even a
     rough quarter or fiscal year? This is a simple ask; never leave the
     completion date as a "[TBD]" without having asked.
   - Quantify the status quo: put a rough number on today's pain — current cost,
     support volume, manual hours, risk exposure. A baseline pre-empts "how do
     we know this is a real problem?"
   - Contention with in-flight work: does this draw on the same scarce people as
     a major program already underway? Flag the overlap directionally.
   - Decision authority: who is the single accountable decision-maker? If several
     units share the outcome, flag the likely need for a steering committee.

7. NEXT STEPS: after a brief, reviewers sometimes recommend an intermediate step
   — a business case/ROI analysis or a requirements-gathering effort — before
   committing to a charter. Don't assume every approved brief goes straight to a
   charter; if cost/value or scope is still fuzzy, naming that intermediate step
   is the honest, helpful move.

---

"""


# ---------------------------------------------------------------------------
# Project helpers
# ---------------------------------------------------------------------------

def _pdir(pid: str) -> Path:
    return PROJECTS_DIR / pid


def _safe_pid(pid: str) -> bool:
    return bool(re.fullmatch(r"[a-f0-9]{8,32}", pid)) and _pdir(pid).is_dir()


def _load_meta(pid: str) -> dict:
    f = _pdir(pid) / "meta.json"
    return json.loads(f.read_text()) if f.is_file() else {}


def _save_meta(pid: str, meta: dict):
    (_pdir(pid) / "meta.json").write_text(json.dumps(meta, indent=2))


def _load_messages(pid: str) -> list:
    f = _pdir(pid) / "messages.json"
    return json.loads(f.read_text()) if f.is_file() else []


def _save_messages(pid: str, msgs: list):
    (_pdir(pid) / "messages.json").write_text(json.dumps(msgs, indent=2))


def _write_interview_log(pid: str):
    """Mirror the full transcript into out/interview-log.md so a turn digest
    survives even when the live CLI --continue session is gone (server restart,
    reopened old project). The retry-fresh path points the model at this file, so
    interview answers that never made it into a section aren't lost on recovery.
    Regenerated wholesale from messages.json each turn — idempotent, always
    complete, no incremental-append drift."""
    msgs = _load_messages(pid)
    if not msgs:
        return
    lines = [
        "# Interview log",
        "",
        "_Verbatim conversation so far. If the live session was lost, read this to",
        "recover the author's answers (including any not yet folded into a section)",
        "before continuing — do not re-ask what they already told you._",
        "",
    ]
    for m in msgs:
        who = "Author" if m.get("role") == "user" else "Assistant"
        lines += [f"## {who}", (m.get("text") or "").strip(), ""]
    out = _pdir(pid) / "out"
    out.mkdir(exist_ok=True)
    (out / "interview-log.md").write_text("\n".join(lines))


def _link_skill(skdir: Path, name: str, src: Path):
    dst = skdir / name
    if dst.exists() or dst.is_symlink():
        return
    try:
        dst.symlink_to(src.resolve(), target_is_directory=True)
    except OSError:
        try:
            shutil.copytree(src, dst)
        except OSError:
            pass


def _ensure_skill_symlink(pid: str):
    """Make the uit-portfolio skill (and the Anthropic docx/pdf/pptx/xlsx helper
    skills) resolvable when Claude Code runs with cwd = the project directory, so
    the model can read uploads and fill .docx via the official docx skill."""
    skdir = _pdir(pid) / ".claude" / "skills"
    skdir.mkdir(parents=True, exist_ok=True)
    _link_skill(skdir, "uit-portfolio", SKILL_DIR)
    for name in HELPER_SKILLS:
        src = GLOBAL_SKILLS_DIR / name
        if src.is_dir():
            _link_skill(skdir, name, src)


def _uploads(pid: str) -> list[str]:
    up = _pdir(pid) / "uploads"
    return sorted(p.name for p in up.iterdir()) if up.is_dir() else []


# Section headings unique to the official Concept Brief template — used to tell a
# real brief apart from any other commented .docx (e.g. a strategy paper).
_BRIEF_MARKERS = [
    "CONCEPT TITLE", "CONCEPT DESCRIPTION", "STRATEGIC ALIGNMENT",
    "OUTCOMES OR KEY RESULTS", "DEFINE SUCCESS", "PROPOSED SOLUTION",
    "PRIORITIZATION DRIVERS", "APPROVED FOR CHARTER",
]


def _is_concept_brief(path: Path) -> bool:
    """True only if the uploaded .docx is (a copy of) the Concept Brief template,
    detected by its characteristic section headings — not just any commented doc."""
    if not str(path).lower().endswith(".docx"):
        return False
    try:
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    parts.append(c.text)
        text = "\n".join(parts).upper()
        hits = sum(1 for m in _BRIEF_MARKERS if m in text)
        return hits >= 3
    except Exception:
        return False


def _latest_doc(pid: str) -> Path | None:
    out = _pdir(pid) / "out"
    docs = sorted(out.glob("*.docx"), key=lambda p: p.stat().st_mtime) if out.is_dir() else []
    return docs[-1] if docs else None


def _content_map(pid: str) -> dict:
    f = _pdir(pid) / "out" / "content.json"
    try:
        return json.loads(f.read_text())
    except Exception:
        return {}


def _project_summary(pid: str) -> dict:
    meta = _load_meta(pid)
    doc = _latest_doc(pid)
    return {
        "id": pid,
        "title": meta.get("title", "Untitled"),
        "mode": meta.get("mode", "brief"),
        "version": meta.get("version", 0),
        "updated": meta.get("updated", 0),
        "has_doc": doc is not None,
        "doc_name": doc.name if doc else None,
    }


# Matches the visible "[TBD — confirm with <who>]" placeholders the skill leaves
# in the content map. Used to surface every open item a reviewer would flag.
_TBD_RE = re.compile(r"\[TBD[^\]]*\]")


def _open_items(c: dict) -> list[dict]:
    """Every unresolved [TBD ...] placeholder in the content map, with the
    section it sits in — this is the "a reviewer will ask about these" checklist
    the doc pane renders. Derived purely from content.json; no skill change."""
    label_map = dict(FIELD_LABELS + SECTION_LABELS)
    label_map["concept_title"] = "Concept Title"
    items = []
    for k in ["concept_title", *(k for k, _ in FIELD_LABELS), *(k for k, _ in SECTION_LABELS)]:
        val = str(c.get(k, ""))
        for m in _TBD_RE.findall(val):
            items.append({"section": label_map.get(k, k), "text": m})
    return items


def build_preview(pid: str) -> dict:
    """Structured preview for the document pane, from out/content.json."""
    c = _content_map(pid)
    meta = _load_meta(pid)
    fields = [{"key": k, "label": lbl, "value": c.get(k, "").strip()}
              for k, lbl in FIELD_LABELS if c.get(k, "").strip()]
    sections = [{"key": k, "label": lbl, "value": c.get(k, "").strip()}
                for k, lbl in SECTION_LABELS if c.get(k, "").strip()]
    return {
        "mode": meta.get("mode", "brief"),
        "title": c.get("concept_title", "").strip() or meta.get("title", "Untitled"),
        "fields": fields,
        "sections": sections,
        "open_items": _open_items(c),
        "version": meta.get("version", 0),
        "empty": not (fields or sections),
    }


# ---------------------------------------------------------------------------
# Claude invocation
# ---------------------------------------------------------------------------

def get_claude_bin() -> str | None:
    """Path to the claude CLI, or None when unavailable / inside Claude Code."""
    if os.environ.get("CLAUDECODE"):
        return None  # nested invocation crashes — fall back to demo
    for d in [".local/bin", "bin"]:
        cand = Path.home() / d / "claude"
        if cand.is_file():
            return str(cand)
    return shutil.which("claude")


def _sse(event: str, data: str) -> str:
    escaped = data.replace("\n", "\ndata: ")
    return f"event: {event}\ndata: {escaped}\n\n"


def _trace(t: str, **kw) -> str:
    """A raw-activity event for the live "Claude Code console" popup."""
    return _sse("trace", json.dumps({"t": t, **kw}))


def _tool_detail(inp: dict) -> str:
    """Best single-line detail for a tool call (for the raw console)."""
    d = (inp.get("file_path") or inp.get("path") or inp.get("command")
         or inp.get("pattern") or inp.get("query") or inp.get("url") or "")
    if not d and inp:
        try:
            d = json.dumps(inp)
        except Exception:
            d = ""
    return str(d)[:240]


def _humanize_tool(tool: str, inp: dict) -> str:
    """Turn a raw tool call into a friendly activity line."""
    detail = inp.get("file_path") or inp.get("path") or inp.get("command") or inp.get("pattern") or ""
    name = Path(detail).name if detail else ""
    friendly = {
        "Read": f"Reading {name}" if name else "Reading your upload",
        "Write": f"Writing {name}" if name else "Drafting the document",
        "Edit": f"Updating {name}" if name else "Revising a section",
        "Bash": "Generating the .docx" if "fill_brief" in str(detail) else "Running a step",
        "Glob": "Looking through your files",
        "Grep": "Scanning for reviewer comments",
    }
    return friendly.get(tool, f"{tool}…")


def _ensure_1m(model: str) -> str:
    """Request the 1M-token context window (``[1m]`` suffix) for the Opus and
    Fable families, which support it — whatever form the id arrived in (short
    first-party alias or full Bedrock inference-profile). Idempotent, and a no-op
    for models that don't take the suffix (e.g. Sonnet/Haiku) or that already
    carry it. The CLI honors ``[1m]`` on either the alias or the profile id."""
    low = (model or "").lower()
    if ("opus" in low or "fable" in low) and "[1m]" not in low:
        return model + "[1m]"
    return model


def resolve_model() -> str:
    """The model id to pass to `claude --model`, Bedrock-aware.

    On Bedrock the identifier must be a full inference-profile id
    (``global.anthropic.claude-opus-4-8``) — the first-party short alias
    (``claude-opus-4-8``) is rejected with "model identifier is invalid". So:
      1. An explicit UIT_CLAUDE_MODEL always wins.
      2. Else, when pointed at Bedrock (CLAUDE_CODE_USE_BEDROCK), use the env's
         ANTHROPIC_MODEL / ANTHROPIC_DEFAULT_OPUS_MODEL (already a full profile id),
         falling back to the known global Opus profile.
      3. Else (first-party API), the short alias is correct.
    Whatever the source, the Opus/Fable families are pinned to the ``[1m]``
    (1M-token context) variant via _ensure_1m.
    """
    explicit = os.environ.get("UIT_CLAUDE_MODEL")
    if explicit:
        return _ensure_1m(explicit)
    if os.environ.get("CLAUDE_CODE_USE_BEDROCK"):
        return _ensure_1m(
            os.environ.get("ANTHROPIC_MODEL")
            or os.environ.get("ANTHROPIC_DEFAULT_OPUS_MODEL")
            or "global.anthropic.claude-opus-4-8[1m]")
    return _ensure_1m("claude-opus-4-8")


# USD per million tokens: (input, output). Cache write bills at 1.25x input,
# cache read at 0.1x input. This app runs Opus by default; the tier is inferred
# from the resolved model id so the estimate tracks whatever model is configured.
_TIER_PRICING_PER_MTOK = {
    "opus": (5.0, 25.0),
    "fable": (10.0, 50.0),
    "sonnet": (3.0, 15.0),
    "haiku": (1.0, 5.0),
}


def _price_tier(model: str) -> str:
    m = (model or "").lower()
    for tier in ("opus", "fable", "sonnet", "haiku"):
        if tier in m:
            return tier
    return "opus"  # this app's default model family


class _CostTracker:
    """Running cost estimate for one CLI invocation.

    stream-json ``assistant`` events carry ``message.usage``; the same message id
    can reappear across events with growing usage, so usage is stored per id
    (latest snapshot wins) and summed — never double-counted. The final ``result``
    event's ``total_cost_usd`` replaces the estimate with the CLI's exact figure."""

    def __init__(self, model: str):
        in_mtok, out_mtok = _TIER_PRICING_PER_MTOK[_price_tier(model)]
        self._in = in_mtok / 1_000_000
        self._out = out_mtok / 1_000_000
        self._usage_by_msg: dict[str, dict] = {}
        self.total_usd = 0.0

    def add_usage(self, msg_id: str, usage: dict) -> None:
        self._usage_by_msg[msg_id] = usage
        total = 0.0
        for u in self._usage_by_msg.values():
            # `... or 0` guards against an explicit null in the stream (get's
            # default only covers a missing key, not a present-but-None value).
            total += (u.get("input_tokens") or 0) * self._in
            total += (u.get("output_tokens") or 0) * self._out
            total += (u.get("cache_creation_input_tokens") or 0) * self._in * 1.25
            total += (u.get("cache_read_input_tokens") or 0) * self._in * 0.1
        self.total_usd = total

    def set_exact(self, total_usd: float) -> None:
        self.total_usd = total_usd


async def stream_claude_cli(pid: str, prompt: str, first_turn: bool,
                            _retrying: bool = False):
    """Run the Claude CLI in the project dir, streaming humanized events."""
    claude_bin = get_claude_bin()
    proj = _pdir(pid)
    # Opus with the 1M-token context window by default, for both turns.
    model = resolve_model()
    effort = os.environ.get("UIT_CLAUDE_EFFORT", "high")  # low|medium|high|xhigh|max
    cmd = [claude_bin, "-p", prompt, "--model", model, "--effort", effort,
           "--output-format", "stream-json", "--verbose",
           "--dangerously-skip-permissions"]
    if not first_turn:
        cmd.insert(1, "--continue")

    env = os.environ.copy()
    home = Path.home()
    prepend = [str(home / "bin"), str(home / ".local" / "bin")]
    parts = env.get("PATH", "").split(":")
    env["PATH"] = ":".join([p for p in prepend if p not in parts] + parts)
    env["CLAUDE_EFFORT"] = effort  # be authoritative — don't inherit a stray xhigh

    proc = await asyncio.create_subprocess_exec(
        *cmd, stdin=asyncio.subprocess.DEVNULL,
        stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE,
        cwd=str(proj), env=env,
    )
    # Drain stderr concurrently — an undrained PIPE deadlocks the CLI once it
    # writes more than the pipe buffer (~64KB) of warnings/progress noise.
    stderr_task = asyncio.ensure_future(proc.stderr.read())

    cost = _CostTracker(model)
    streamed_text = ""  # everything the user saw streamed live
    result_text = ""    # the final `result` event (usually a duplicate of the last block)
    buf = ""
    try:
        yield _trace("meta", text=f"$ claude --model {model} --effort {effort}"
                                  + (" --continue" if not first_turn else ""))
        while True:
            chunk = await asyncio.wait_for(proc.stdout.read(512), timeout=600)
            if not chunk:
                break
            buf += chunk.decode("utf-8", errors="replace")
            while "\n" in buf:
                line, buf = buf.split("\n", 1)
                line = line.strip()
                if not line:
                    continue
                try:
                    event = json.loads(line)
                except json.JSONDecodeError:
                    continue
                etype = event.get("type")
                if etype == "assistant":
                    msg = event.get("message", {})
                    usage = msg.get("usage")
                    if isinstance(usage, dict):
                        cost.add_usage(msg.get("id", ""), usage)
                        yield _sse("cost", f"{cost.total_usd:.4f}")
                    for block in event.get("message", {}).get("content", []):
                        if block.get("type") == "text" and block.get("text"):
                            streamed_text += block["text"]
                            yield _sse("chunk", block["text"])
                            yield _trace("text", text=block["text"])
                        elif block.get("type") == "tool_use":
                            inp = block.get("input", {})
                            yield _sse("activity", _humanize_tool(block.get("name", ""), inp))
                            yield _trace("tool", name=block.get("name", ""), detail=_tool_detail(inp))
                elif etype == "user":
                    for block in event.get("message", {}).get("content", []):
                        if block.get("type") != "tool_result":
                            continue
                        raw = block.get("content", "") or block.get("output", "")
                        if isinstance(raw, list):
                            raw = "\n".join(b.get("text", "") for b in raw
                                            if isinstance(b, dict) and b.get("type") == "text")
                        if raw and isinstance(raw, str):
                            lines = raw.strip().splitlines()
                            preview = lines[0][:200] if lines else ""
                            suffix = f" …(+{len(lines) - 1} lines)" if len(lines) > 1 else ""
                            yield _trace("result", detail=preview + suffix)
                elif etype == "result":
                    exact = event.get("total_cost_usd")
                    if isinstance(exact, (int, float)) and exact > 0:
                        cost.set_exact(float(exact))
                        yield _sse("cost", f"{cost.total_usd:.4f}")
                    rt = event.get("result", "")
                    if rt and isinstance(rt, str):
                        result_text = rt
    except asyncio.TimeoutError:
        proc.kill()
        stderr_task.cancel()
        yield _trace("meta", text="✕ timed out after 600s")
        yield _sse("error", "Claude timed out.")
        return
    except asyncio.CancelledError:
        # Client disconnected mid-stream — don't orphan the CLI process or
        # leak the stderr drain task.
        proc.kill()
        stderr_task.cancel()
        raise

    await proc.wait()
    if proc.returncode != 0:
        try:
            # Bounded wait: a grandchild that inherited the stderr pipe could
            # keep it open past our process's exit and hang this await forever.
            err = (await asyncio.wait_for(stderr_task, timeout=5)).decode(
                "utf-8", errors="replace")
        except (asyncio.TimeoutError, asyncio.CancelledError):
            err = ""
        # A --continue against a session that no longer exists (server restart,
        # reopened old project) fails before producing any output. Fall back to
        # a fresh session once, re-grounding the model in the on-disk state.
        if not first_turn and not _retrying and not streamed_text.strip():
            yield _trace("meta", text=f"✕ exit {proc.returncode} — no session to "
                                      "continue; retrying fresh")
            fresh = (PREAMBLE
                     + "[Context: this resumes an existing project, but the prior "
                       "CLI session is gone (e.g. after a server restart). Before "
                       "responding, re-read ./out/content.json for the current "
                       "document state, ./out/interview-log.md for the full "
                       "conversation and the author's answers so far (don't re-ask "
                       "what they already told you), and ./uploads/ for their "
                       "files.]\n\n"
                     + prompt)
            async for ev in stream_claude_cli(pid, fresh, first_turn=True,
                                              _retrying=True):
                yield ev
            return
        yield _trace("meta", text=f"✕ exit {proc.returncode}")
        yield _sse("error", f"Claude exited {proc.returncode}: {err[:400]}")
        return
    stderr_task.cancel()
    yield _trace("meta", text="✓ completed")

    # The text was already streamed live; persist it but do NOT re-emit it.
    transcript = streamed_text.strip() or result_text.strip()
    async for ev in _finish_turn(pid, transcript, first_turn, emit_final=False):
        yield ev


# --- Demo mode -------------------------------------------------------------

_DEMO_TODAY = time.localtime()
_DEMO_BASE = {
    "concept_title": "Enterprise Digital Signage Platform",
    "requestor": "[your name]",
    "requestor_title_unit": "University IT",
    "sponsor": "[TBD — confirm sponsor]",
    "request_date": f"{time.strftime('%b', _DEMO_TODAY)} {_DEMO_TODAY.tm_mday}, {_DEMO_TODAY.tm_year}",
    "concept_description": ("Each college currently runs its own digital-signage hardware with no "
                            "central control, so security patches lag for months. This proposes a "
                            "single enterprise platform managed by one team."),
    "strategic_alignment": "[TBD — name the specific Prosperity Widely Shared commitment this advances]",
    "outcomes": ("- One centrally managed, regularly patched platform replaces per-college systems\n"
                 "- Clear single ownership for updates and security posture"),
    "benefits": "Without this, signage endpoints stay months behind on security patches — a standing risk surface.",
    "risks": "Per-college boxes run months behind on patches; migration effort across units is unscoped.",
    "define_success": "[TBD — an observable end-state, e.g. 100% of in-scope displays on the managed platform]",
    "proposed_solution": "A directional approach: consolidate onto one managed platform with central patching. (Not a final design.)",
    "prioritization_drivers": "Security exposure from unpatched endpoints.",
    "key_requirements": "Central management console; automated patching; role-based access per college.",
    "teams_and_skills": "UIT platform team; distributed IT in each college; OIS for security review.",
    "additional_notes": "",
}

_DEMO_SCRIPT_BRIEF = [
    ("activity", "Reading your idea and mapping it onto the brief template…"),
    ("chunk", "Good — I’ve mapped what you gave me onto the Concept Brief. **Concept Description** and **Risks** are in reasonable shape.\n\n"),
    ("activity", "Checking against the Tier 1 reviewer questions…"),
    ("chunk", "Before a reviewer sees this, a few directional things are worth pinning down — these are the lightweight Tier 1 questions; the deeper operational detail belongs in the charter, not here.\n\n"),
    ("chunk", "```questions\nWhat’s driving the timing — a security finding, an audit, or leadership direction?\nRoughly how many displays are affected today?\nHow far behind are patches on the current systems?\nIs this adding a new platform, or replacing existing per-college systems?\nWho is the single accountable decision-maker?\n```\n"),
    ("chunk", "Answer what you can — “I don’t know” is fine and I’ll mark it `[TBD]`. I’ve already drafted enough for a solid first version you can react to.\n\n"),
    ("activity", "Generating the .docx in the official template…"),
]

_DEMO_SCRIPT_CHARTER = [
    ("activity", "Reading the uploaded brief and reviewer comments…"),
    ("chunk", "I can see the reviewer comments. I’ll treat each one as a required input and make sure every question is visibly answered in the charter.\n\n"),
    ("activity", "Checking Tier 2 charter questions against the comments…"),
    ("chunk", "The areas reviewers probe hardest at charter stage are scope, data classification, and funding shape. Let’s nail these down:\n\n"),
    ("chunk", "```questions\nWhat is explicitly out of scope for this work?\nWhich data classification or regime governs the data — FERPA, HIPAA, GLBA, PCI, or CUI?\nIs the implementation funding one-time (capital) or recurring (cloud operating)?\nWhich funding year or years does this draw on?\nWho owns technical management and data stewardship after go-live?\n```\n"),
    ("chunk", "Once I have these I’ll fold them in; anything unknown becomes a `[TBD]` placeholder.\n\n"),
    ("activity", "Generating the charter .docx…"),
]


async def stream_demo(pid: str, user_text: str, first_turn: bool):
    """Scripted interview + real .docx generation, for running without a backend."""
    yield _sse("status", "Demo mode — no Claude backend configured; scripted response.")
    yield _trace("meta", text="DEMO MODE — scripted assistant, no Claude backend")
    meta = _load_meta(pid)
    mode = meta.get("mode", "brief")
    script = _DEMO_SCRIPT_CHARTER if mode == "charter" else _DEMO_SCRIPT_BRIEF
    for kind, text in script:
        await asyncio.sleep(0.12)
        yield _sse(kind, text)
        if kind == "activity":
            yield _trace("tool", name="step", detail=text)
        else:
            yield _trace("text", text=text)

    # Build/refresh the content map from any prior state, appending the author's note.
    content = dict(_content_map(pid)) or dict(_DEMO_BASE)
    if user_text.strip():
        note = user_text.strip()
        existing = content.get("additional_notes", "").strip()
        content["additional_notes"] = (existing + "\n" if existing else "") + f"- Author note: {note}"
    (_pdir(pid) / "out").mkdir(exist_ok=True)
    (_pdir(pid) / "out" / "content.json").write_text(json.dumps(content, indent=2))

    await asyncio.sleep(0.15)
    open_tbds = [k for k, v in content.items() if "[TBD" in str(v)]
    tbd_note = (f" Open placeholders: {', '.join(open_tbds)}." if open_tbds else "")
    final_text = (
        "I've drafted the document — check the preview on the right and download when ready."
        + tbd_note
        + "\n\nTell me anything you want to tighten and I'll revise."
    )
    async for ev in _finish_turn(pid, final_text, first_turn, demo=True):
        yield ev


# --- shared turn finalization ---------------------------------------------

async def _finish_turn(pid: str, final_text: str, first_turn: bool, demo: bool = False,
                       emit_final: bool = True):
    """After the model (or demo) ran: (re)generate the .docx, bump version,
    persist transcript, emit doc + preview + done events.

    emit_final: when True, stream final_text as a chunk (demo mode, where the
    closing line was not otherwise streamed). The CLI path passes False because
    the text was already streamed live — re-emitting would duplicate it."""
    meta = _load_meta(pid)

    # In demo mode we own generation. In CLI mode the skill already wrote the
    # docx; if it somehow didn't (and we have a content map), generate as a
    # safety net so the download/preview always work.
    out = _pdir(pid) / "out"
    out.mkdir(exist_ok=True)
    need_gen = demo or (_latest_doc(pid) is None and (out / "content.json").is_file())
    mode = meta.get("mode", "brief")
    if need_gen and (out / "content.json").is_file() and mode == "brief":
        # Hard guard: never fill the Word template unless the docx skill is loaded.
        if not docx_skill_available():
            yield _trace("meta", text="✕ docx skill not installed — template not filled")
            yield _sse("error", "The Anthropic **docx** skill isn't installed "
                                "(~/.claude/skills/docx), so the Word document was not "
                                "generated. Install the docx skill and try again.")
        else:
            ver = meta.get("version", 0) + 1
            title = _content_map(pid).get("concept_title") or meta.get("title", "Concept")
            safe = re.sub(r"[^A-Za-z0-9]+", "_", title).strip("_")[:40] or "Concept"
            out_name = f"{safe}_Concept_Brief_v{ver}.docx"
            try:
                subprocess.run(
                    ["python3", str(FILL_BRIEF), "--template", str(BRIEF_TEMPLATE),
                     "--out", str(out / out_name), "--content", str(out / "content.json")],
                    check=True, capture_output=True, text=True, timeout=60,
                )
            except subprocess.CalledProcessError as e:
                yield _sse("error", f"Document generation failed: {e.stderr[:300]}")
    # Charter generation is handled by the skill (docx skill / python-docx); no fallback script.

    # Promote the model's concise concept_title to the project title once it
    # exists, replacing the provisional "New …" placeholder — a far better label
    # than the author's first sentence.
    ctitle = _content_map(pid).get("concept_title", "").strip()
    if ctitle and not ctitle.startswith("[TBD") and meta.get("title", "").startswith(("New ", "Untitled")):
        meta["title"] = ctitle[:60]

    # Bump the version whenever the document file is newer than what we last
    # recorded. This works for both demo mode (which writes a new _vN file) and
    # the CLI path (where the skill overwrites the same filename each turn) —
    # comparing mtimes catches a regenerated doc that kept its name. A turn that
    # only asks questions and doesn't touch the doc leaves the version unchanged.
    doc = _latest_doc(pid)
    if doc:
        mt = doc.stat().st_mtime
        if mt > meta.get("doc_mtime", 0):
            meta["version"] = meta.get("version", 0) + 1
            meta["doc_mtime"] = mt
    if final_text.strip():
        msgs = _load_messages(pid)
        msgs.append({"role": "assistant", "text": final_text, "ts": time.time()})
        _save_messages(pid, msgs)
    meta["updated"] = time.time()
    _save_meta(pid, meta)
    # Persist a recoverable digest of the whole conversation (item 6).
    _write_interview_log(pid)

    if final_text.strip() and emit_final:
        yield _sse("chunk", ("\n\n" if not final_text.startswith("\n") else "") + final_text)
    if doc:
        yield _sse("doc", json.dumps({
            "name": doc.name,
            "url": f"/api/projects/{pid}/download/{doc.name}",
            "version": meta.get("version", 0),
        }))
    yield _sse("preview", json.dumps(build_preview(pid)))
    yield _sse("done", "")


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

# Projects with a turn in flight. A single Claude Code process per project owns
# its dir (the CLI session, out/, content.json); a second concurrent turn would
# race --continue and clobber files. asyncio is single-threaded, so claiming the
# pid synchronously (no await between the check and the add) is an atomic guard
# against a double-submit or a second browser tab.
_active_turns: set[str] = set()


@app.get("/api/version")
async def version():
    return JSONResponse({
        "commit": _GIT_COMMIT,
        "demo": get_claude_bin() is None,
        "docx_skill": docx_skill_available(),
        "missing_skills": missing_skills(),
        "contact": os.environ.get("CONTACT", ""),
        "coordinator": os.environ.get("PORTFOLIO_COORDINATOR", ""),
    })


@app.get("/api/projects")
async def list_projects():
    out = []
    for d in PROJECTS_DIR.iterdir():
        if d.is_dir() and (d / "meta.json").is_file():
            out.append(_project_summary(d.name))
    out.sort(key=lambda p: p["updated"], reverse=True)
    return JSONResponse(out)


@app.post("/api/projects")
async def create_project(request: Request):
    body = await request.json()
    mode = body.get("mode", "brief")
    if mode not in ("brief", "charter"):
        mode = "brief"
    title = (body.get("title") or "").strip() or ("New Charter" if mode == "charter" else "New Concept Brief")
    pid = uuid.uuid4().hex[:16]
    p = _pdir(pid)
    (p / "uploads").mkdir(parents=True)
    (p / "out").mkdir(parents=True)
    _ensure_skill_symlink(pid)
    _save_meta(pid, {"title": title, "mode": mode, "version": 0,
                     "created": time.time(), "updated": time.time()})
    _save_messages(pid, [])
    return JSONResponse(_project_summary(pid))


@app.get("/api/projects/{pid}")
async def get_project(pid: str):
    if not _safe_pid(pid):
        return JSONResponse({"error": "not found"}, status_code=404)
    return JSONResponse({
        **_project_summary(pid),
        "messages": _load_messages(pid),
        "uploads": _uploads(pid),
        "preview": build_preview(pid),
    })


@app.delete("/api/projects/{pid}")
async def delete_project(pid: str):
    if _safe_pid(pid):
        shutil.rmtree(_pdir(pid), ignore_errors=True)
    return JSONResponse({"ok": True})


@app.post("/api/projects/{pid}/upload")
async def upload(pid: str, file: UploadFile = File(...)):
    if not _safe_pid(pid):
        return JSONResponse({"error": "not found"}, status_code=404)
    safe_name = Path(file.filename or "upload").name
    dest = _pdir(pid) / "uploads" / safe_name
    data = await file.read()
    dest.write_bytes(data)

    # Detect Word reviewer comments (a charter signal), AND whether this is
    # actually a Concept Brief — the charter choice is only offered for a brief.
    has_comments = False
    if safe_name.lower().endswith(".docx"):
        try:
            import zipfile
            z = zipfile.ZipFile(dest)
            has_comments = ("word/comments.xml" in z.namelist()
                            and bool(z.read("word/comments.xml").strip()))
        except Exception:
            pass
    is_brief = _is_concept_brief(dest)
    meta = _load_meta(pid)
    meta["updated"] = time.time()
    _save_meta(pid, meta)
    # NEVER auto-switch to charter here. We only report is_brief/has_comments so the
    # UI can ASK — the mode changes only via an explicit choice (set_mode).
    # `suggested_mode` is the UI's *pre-selected* default for that ask, not a switch:
    # a Concept Brief that came back WITH reviewer comments is one ready to charter,
    # so charter is the default; a brief with no comments stays ambiguous (brief).
    suggested_mode = "charter" if (is_brief and has_comments) else "brief"
    return JSONResponse({"uploads": _uploads(pid), "has_comments": has_comments,
                         "is_brief": is_brief, "mode": meta.get("mode", "brief"),
                         "suggested_mode": suggested_mode})


@app.post("/api/projects/{pid}/mode")
async def set_mode(pid: str, request: Request):
    """Change a project's mode — only ever called from an explicit user choice."""
    if not _safe_pid(pid):
        return JSONResponse({"error": "not found"}, status_code=404)
    body = await request.json()
    mode = body.get("mode")
    if mode not in ("brief", "charter"):
        return JSONResponse({"error": "invalid mode"}, status_code=400)
    meta = _load_meta(pid)
    meta["mode"] = mode
    meta["updated"] = time.time()
    _save_meta(pid, meta)
    return JSONResponse(_project_summary(pid))


@app.post("/api/projects/{pid}/message")
async def message(pid: str, request: Request):
    if not _safe_pid(pid):
        return JSONResponse({"error": "not found"}, status_code=404)
    # Claim the project synchronously (before any await) so a double-submit or a
    # second tab can't start a second turn racing the same CLI session / out dir.
    if pid in _active_turns:
        return JSONResponse(
            {"error": "A response is already being generated for this project. "
                      "Wait for it to finish before sending another message."},
            status_code=409,
        )
    _active_turns.add(pid)
    try:
        body = await request.json()
        text = (body.get("text") or "").strip()

        msgs = _load_messages(pid)
        first_turn = not any(m["role"] == "assistant" for m in msgs)
        if text:
            msgs.append({"role": "user", "text": text, "ts": time.time()})
            _save_messages(pid, msgs)

        meta = _load_meta(pid)
        # NOTE: the project title is intentionally NOT set from the raw first
        # message (that produced an ugly sentence-fragment label). It is promoted
        # from the model's concise concept_title in _finish_turn once drafted.

        _ensure_skill_symlink(pid)
    except BaseException:
        # BaseException, not Exception: a client disconnect during `await
        # request.json()` raises asyncio.CancelledError (a BaseException), which
        # would otherwise skip the release and lock the project permanently.
        _active_turns.discard(pid)  # release if setup failed before streaming
        raise

    async def generate():
        try:
            yield _sse("project", json.dumps(_project_summary(pid)))
            if get_claude_bin():
                mode = meta.get("mode", "brief")
                uploads = _uploads(pid)
                ctx = (f"[Project mode: {mode}. Uploaded files: {', '.join(uploads) or 'none'}.]\n\n"
                       f"{text}")
                prompt = (PREAMBLE + ctx) if first_turn else ctx
                async for ev in stream_claude_cli(pid, prompt, first_turn):
                    yield ev
            else:
                async for ev in stream_demo(pid, text, first_turn):
                    yield ev
        finally:
            _active_turns.discard(pid)  # release when the stream ends or is cancelled

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.get("/api/projects/{pid}/preview")
async def preview(pid: str):
    if not _safe_pid(pid):
        return JSONResponse({"error": "not found"}, status_code=404)
    return JSONResponse(build_preview(pid))


@app.get("/api/projects/{pid}/download/{name}")
async def download(pid: str, name: str):
    if not _safe_pid(pid):
        return JSONResponse({"error": "not found"}, status_code=404)
    safe = Path(name).name
    f = _pdir(pid) / "out" / safe
    if not f.is_file():
        return JSONResponse({"error": "not found"}, status_code=404)
    return FileResponse(
        str(f), filename=safe,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/", response_class=HTMLResponse)
async def index():
    # no-store so the single-file SPA is never served stale (the whole app —
    # HTML, CSS, JS — lives in this one document, so a cached copy means old code).
    return HTMLResponse(
        (REPO_ROOT / "static" / "index.html").read_text(encoding="utf-8"),
        headers={"Cache-Control": "no-store, must-revalidate", "Pragma": "no-cache"},
    )


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    reload = not os.environ.get("SYSTEMD_EXEC_PID")
    uvicorn.run("app:app", host="127.0.0.1", port=port, reload=reload)
